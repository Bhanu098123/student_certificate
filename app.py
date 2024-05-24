from flask import Flask, render_template, request, send_file, Response
import pandas as pd
from docx import Document
from docx.shared import Pt
import pythoncom
from docx2pdf import convert
import io
import os
import re

app = Flask(__name__)

# Load student data from Excel
students_df = pd.read_excel('students.xlsx', dtype={'Roll Number': str})

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_certificate', methods=['POST'])
def generate_certificate():
    roll_number = request.form.get('roll_number')

    if not roll_number:
        return "Roll number not provided", 400

    # Convert roll number to string to ensure consistency
    roll_number = str(roll_number)

    try:
        student = students_df[students_df['Roll Number'] == roll_number]
    except KeyError as e:
        return f"Column not found: {e}", 500

    if student.empty:
        return render_template('not_found.html')

    student_details = student.iloc[0]

    # Load the template
    doc = Document('certificate_template.docx')

    # Function to replace placeholders in text
    def replace_placeholders(text, placeholders):
        for key, value in placeholders.items():
            text = text.replace(key, str(value))  # Ensure value is converted to string
        return text

    # Function to replace placeholders in a paragraph
    def replace_placeholders_in_paragraph(paragraph, placeholders):
        combined_text = "".join([run.text for run in paragraph.runs])
        replaced_text = replace_placeholders(combined_text, placeholders)

        # Clear the existing runs
        for run in paragraph.runs:
            run.text = ""

        # Redistribute the replaced text into the runs
        if len(paragraph.runs) > 0:
            paragraph.runs[0].text = replaced_text
            for run in paragraph.runs:
                if replaced_text.strip() in run.text.strip():  # Only bold the replaced text
                    run.bold = True
        else:
            paragraph.add_run(replaced_text)

    # Placeholders dictionary
    placeholders = {
        '{{ student_name }}': student_details['Student Name'],
        '{{ roll_number }}': student_details['Roll Number'],
        '{{ branch_name }}': student_details['Branch Name'],
        '{{ college_name }}': student_details['College Name'],
        '{{ university_name }}': student_details['University Name'],
        '{{ domain_name }}': student_details['Domain Name'],
        '{{ course }}': student_details['Course']
    }

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, placeholders)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, placeholders)

    # Save the modified document to a temporary file
    temp_docx_path = 'temp_certificate.docx'
    temp_pdf_path = 'temp_certificate.pdf'
    doc.save(temp_docx_path)

    # Initialize COM library
    pythoncom.CoInitialize()
    try:
        # Convert the DOCX to PDF
        convert(temp_docx_path, temp_pdf_path)
    finally:
        # Uninitialize COM library
        pythoncom.CoUninitialize()

    # Send the PDF file with roll number as the file name
    with open(temp_pdf_path, 'rb') as f:
        pdf_data = f.read()

    # Clean up temporary files
    os.remove(temp_docx_path)
    os.remove(temp_pdf_path)

    # Set the filename as the roll number
    filename = f'{roll_number}.pdf'

    # Create a response with the PDF data
    response = Response(pdf_data, content_type='application/pdf')
    response.headers.set('Content-Disposition', f'attachment; filename="{filename}"')

    return response

if __name__ == '__main__':
    app.run(debug=True)
